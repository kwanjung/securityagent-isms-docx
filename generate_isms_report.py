#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys, os, re
from datetime import datetime
import pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────────────────────
#  공통 상수
# ──────────────────────────────────────────────────────────────────
FONT = "Malgun Gothic"          # 사용자 환경(Word) 기준. Mac은 "Apple SD Gothic Neo" 권장
FILL_MANUAL = "FFF2CC"          # 〔작성 필요〕 수기 셀 (노랑)
FILL_HEAD   = "1F2937"          # 표 헤더 (다크 네이비)
FILL_LABEL  = "EDF2F7"          # 라벨 셀 (연회색)

KNOWN_RISK_TYPES = [
    "Arbitrary File Upload", "Code Injection", "Command Injection", "Cross Site Scripting",
    "Insecure Direct Object Reference", "Json Web Token Vulnerabilities", "Local File Inclusion",
    "Path Traversal", "Privilege Escalation", "Server Side Request Forgery",
    "Server Side Template Injection", "Sql Injection", "Xml External Entity",
]
RISK_KR = {
    "Arbitrary File Upload": "임의 파일 업로드", "Code Injection": "코드 인젝션",
    "Command Injection": "명령어 인젝션", "Cross Site Scripting": "크로스 사이트 스크립팅(XSS)",
    "Insecure Direct Object Reference": "안전하지 않은 직접 객체 참조(IDOR)",
    "Json Web Token Vulnerabilities": "JWT 취약점", "Local File Inclusion": "로컬 파일 포함(LFI)",
    "Path Traversal": "경로 조작(Path Traversal)", "Privilege Escalation": "권한 상승",
    "Server Side Request Forgery": "서버측 요청 위조(SSRF)",
    "Server Side Template Injection": "서버측 템플릿 인젝션(SSTI)",
    "Sql Injection": "SQL 인젝션", "Xml External Entity": "XML 외부 엔티티(XXE)",
}
# 발견 위험유형 → ISMS-P 인증기준 보조 매핑 (직접 증적 2.11.2 는 항상 포함)
RISK_TO_ISMS = {
    "Path Traversal": [("2.6.3 응용프로그램 접근", "경로 파라미터 인가·접근통제 우회", "보조 증적"),
                       ("2.10.3 공개서버 보안", "대외 공개 서버 대상 점검", "보조 증적"),
                       ("2.7.1 암호정책 적용", "트래버설로 인한 자격증명·비밀 노출", "보조 증적")],
    "Sql Injection": [("2.6.4 데이터베이스 접근", "SQL 인젝션·데이터 노출", "보조 증적")],
    "Cross Site Scripting": [("2.6.3 응용프로그램 접근", "입력검증·출력인코딩 미흡", "보조 증적")],
    "Privilege Escalation": [("2.5.5 특수 계정 및 권한 관리", "권한 상승", "보조 증적"),
                             ("2.6.3 응용프로그램 접근", "인가 우회", "보조 증적")],
    "Insecure Direct Object Reference": [("2.6.3 응용프로그램 접근", "객체 참조 인가통제 미흡(IDOR)", "보조 증적")],
    "Json Web Token Vulnerabilities": [("2.5.3 사용자 인증", "토큰 서명·검증 취약", "보조 증적")],
    "Server Side Request Forgery": [("2.6.7 인터넷 접속 통제", "내부 자원·메타데이터 접근(SSRF)", "보조 증적")],
}

# CVSS 메트릭/판정값 한글 (표준 용어 — API 번역 불필요)
METRIC_KR = {
    "Attack Vector (AV)": "공격 벡터 (AV)", "Attack Complexity (AC)": "공격 복잡도 (AC)",
    "Privileges Required (PR)": "필요 권한 (PR)", "User Interaction (UI)": "사용자 상호작용 (UI)",
    "Scope (S)": "영향 범위 (S)", "Confidentiality Impact (C)": "기밀성 영향 (C)",
    "Integrity Impact (I)": "무결성 영향 (I)", "Availability Impact (A)": "가용성 영향 (A)",
}
VAL_KR = {
    "network": "네트워크(Network)", "adjacent": "인접(Adjacent)", "local": "로컬(Local)",
    "physical": "물리(Physical)", "low": "낮음(Low)", "high": "높음(High)", "none": "없음(None)",
    "required": "필요(Required)", "unchanged": "변경 없음(Unchanged)", "changed": "변경됨(Changed)",
}

# Amazon Bedrock 번역 기본값 — Opus 4.6 글로벌 추론 프로파일(global CRIS).
# 계정/리전에 따라 --model 로 apac.* / us.* / 다른 모델ID 로 교체 가능.
DEFAULT_MODEL_ID = "global.anthropic.claude-opus-4-6-v1"

# ══════════════════════════════════════════════════════════════════
#  1. PDF 추출
# ══════════════════════════════════════════════════════════════════
def _clean(s):
    """하이픈 줄바꿈 결합."""
    return re.sub(r"-\n(\w)", r"\1", s or "")

def _reformat_date(s):
    for fmt in ("%B %d, %Y",):
        try: return datetime.strptime(s.strip(), fmt).strftime("%Y-%m-%d")
        except Exception: pass
    return s

def _reformat_dt(s):
    m = re.search(r"(\d+)/(\d+)/(\d+),\s*(\d+):(\d+):(\d+)\s*(AM|PM)", s or "")
    if not m: return s
    try:
        dt = datetime.strptime(f"{m.group(1)}/{m.group(2)}/{m.group(3)}, "
                               f"{m.group(4)}:{m.group(5)}:{m.group(6)} {m.group(7)}",
                               "%m/%d/%Y, %I:%M:%S %p")
        return dt.strftime("%Y-%m-%d %H:%M:%S (KST, GMT+9)")
    except Exception:
        return s

def extract_report(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        # 각 페이지 말미의 단독 페이지번호(푸터) 제거 후 결합
        page_texts = [re.sub(r"\n\s*\d{1,3}\s*$", "", (p.extract_text() or "").rstrip())
                      for p in pdf.pages]
        full = _clean("\n".join(page_texts))

        # 점검 대상/일자
        m = re.search(r"penetration test of the application hosted at\s+(\S+)\s+on\s+"
                      r"([A-Z][a-z]+ \d+,\s*\d{4})", full)
        target_url = m.group(1) if m else "(추출 실패)"
        tested_on  = _reformat_date(m.group(2)) if m else "(추출 실패)"

        # 심각도 카운트
        m = re.search(r"comprising\s+(\d+)\s+critical-severity,\s+(\d+)\s+high-severity,\s+"
                      r"(\d+)\s+medium-severity,\s+(\d+)\s+low-severity,\s+and\s+"
                      r"(\d+)\s+informational-severity", full)
        counts = dict(zip(("critical", "high", "medium", "low", "info"),
                          map(int, m.groups()))) if m else {}
        total = sum(counts.values()) if counts else None

        # 범위 도메인
        scope_block = full[full.find("\nScope"): full.find("Methodology")] if "\nScope" in full else ""
        domains = sorted(set(re.findall(r"https?://[^\s•]+", scope_block))) or \
                  ([target_url] if target_url.startswith("http") else [])

        # 리포트 필터 고지
        def _filt(label):
            mm = re.search(rf"{label}:\s*\n([^\n]+)", full)
            return mm.group(1).strip() if mm else "(전체)"
        mrt = re.search(r"Risk types:\s*\n(.+?)\nTask status", full, re.S)
        risktypes = re.sub(r"\s+", " ", mrt.group(1)).strip() if mrt else "(전체)"
        filters = {
            "risk":   _filt("Risk levels"), "confidence": _filt("Confidence"),
            "status": _filt("Status"),      "risktypes":  risktypes,
        }

        # 수행 위험유형(좌측 컬럼) + 태스크 총계(상태 컬럼)
        try:
            start = next(i for i, t in enumerate(page_texts) if "Tasks Executed" in t)
            end   = next(i for i, t in enumerate(page_texts) if i > start and "Findings (" in t)
        except StopIteration:
            start, end = 0, len(page_texts)
        left_tokens, comp, abort = [], 0, 0
        for i in range(start, end):
            for w in sorted(pdf.pages[i].extract_words(), key=lambda w: (round(w["top"] / 2), w["x0"])):
                if w["x0"] < 175 and w["text"] != "Task":
                    left_tokens.append(w["text"])
                elif w["x0"] >= 460 and w["text"] == "Completed": comp += 1
                elif w["x0"] >= 460 and w["text"] == "Aborted":   abort += 1
        ls = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", re.sub(r"-\s+", "", " ".join(left_tokens)))
        risk_types = [rt for rt in KNOWN_RISK_TYPES if rt in ls]
        tasks = {"types": risk_types, "total": comp + abort, "completed": comp, "aborted": abort}

        # 발견 상세
        findings = []
        for fm in re.finditer(r"Finding \d+:\s*(.+?)\nID\s+(\S+)", full, re.S):
            blk = full[fm.start(): fm.start() + 6000]
            title = re.sub(r"\s+", " ", fm.group(1)).strip()

            def field(lbl):
                mm = re.search(rf"\n{re.escape(lbl)}\s+(.+)", blk)
                return mm.group(1).strip() if mm else None

            def section(a, b):
                ia = blk.find(a)
                ib = blk.find(b, ia + len(a)) if ia >= 0 else -1
                if ia < 0: return ""
                return re.sub(r"\s+", " ", blk[ia + len(a): ib if ib > 0 else ia + 3000]).strip()

            desc = re.sub(r"\s+\d{1,3}$", "", section("Description", "Reproduction Steps"))
            repro_raw = section("Reproduction Steps", "Risk Reasoning")
            rr_text = re.sub(r"\s+", " ", _clean(blk[blk.find("Risk Reasoning"):])) if "Risk Reasoning" in blk else ""
            rr = []
            for mm in re.finditer(
                    r"-\s*([A-Za-z][A-Za-z ]*\([A-Z]+\)):\s*([A-Za-z]+)\s*-\s*(.+?)"
                    r"(?=\s*-\s*[A-Za-z][A-Za-z ]*\([A-Z]+\):|\Z)", rr_text, re.S):
                metric, short, expl = mm.group(1).strip(), mm.group(2).strip(), mm.group(3).strip()
                rr.append((metric, short, expl))
            findings.append(dict(
                title=title, id=fm.group(2), severity=field("Severity"),
                risktype=field("Risk Type"), confidence=field("Confidence"),
                status=field("Status"), score=field("Risk Score"),
                cvss=field("CVSS v3.1 Vector"), identified=_reformat_dt(field("Identified on")),
                description=desc, reproduction=repro_raw, risk_reasoning=rr[:8],
            ))

    name = re.search(r"Penetration Test Report\s+(\S+)", full)
    target_name = name.group(1) if name else os.path.splitext(os.path.basename(pdf_path))[0]
    return dict(target=target_name, target_url=target_url, tested_on=tested_on,
                counts=counts, total=total, domains=domains, filters=filters,
                tasks=tasks, findings=findings)


# ══════════════════════════════════════════════════════════════════
#  1.5 번역 (Amazon Bedrock Claude) — 기본 ON
# ══════════════════════════════════════════════════════════════════
_TRANSLATE_PROMPT = (
    "너는 정보보안 취약점 점검 보고서 전문 번역가야. 아래 영문 내용을 한국어로 번역해조.\n"
    "규칙:\n"
    "- 명령어, 코드, 파일경로, URL, 페이로드, HTTP 응답/출력값, CVE·CVSS 식별자, "
    "제품·기술 고유명은 번역하지 말고 원문 그대로 유지\n"
    "- 서술형 설명만 자연스러운 한국어 보고서체(문어체)로 번역\n"
    "- 번역 결과 텍스트만 출력하고 머리말·설명·따옴표를 덧붙이지 않는다.\n\n"
    "[원문]\n"
)

def _has_english(text):
    return bool(re.search(r"[A-Za-z]{3,}", text or ""))

class Translator:
    """영문 추출 텍스트를 한국어로 번역. 백엔드=Amazon Bedrock Claude(Converse).
    Bedrock 사용 불가 시 원문을 그대로 유지하고 경고만 출력(생성은 계속)."""

    def __init__(self, enabled=True, model_id=DEFAULT_MODEL_ID, region=None):
        self.enabled = enabled
        self.model_id = model_id
        self.cache = {}
        self.client = None
        self._warned = False
        if not enabled:
            return
        try:
            import boto3
            self.client = (boto3.client("bedrock-runtime", region_name=region)
                           if region else boto3.client("bedrock-runtime"))
        except Exception as e:
            self._warn(f"Bedrock 클라이언트 생성 실패 → 원문(영문) 유지: {e}")
            self.enabled = False

    def _warn(self, msg):
        if not self._warned:
            print(f"[번역 경고] {msg}", file=sys.stderr)
            print("           (번역 없이 영문 원문으로 계속 진행)", file=sys.stderr)
            self._warned = True

    def _invoke(self, text):
        resp = self.client.converse(
            modelId=self.model_id,
            messages=[{"role": "user", "content": [{"text": _TRANSLATE_PROMPT + text}]}],
            inferenceConfig={"maxTokens": 4096, "temperature": 0},
        )
        blocks = resp["output"]["message"]["content"]
        return "".join(b.get("text", "") for b in blocks).strip()

    def __call__(self, text):
        if not self.enabled or not text or not _has_english(text):
            return text
        if text in self.cache:
            return self.cache[text]
        try:
            out = self._invoke(text) or text
        except Exception as e:
            self._warn(f"번역 호출 실패 → 원문(영문) 유지: {e}")
            self.enabled = False
            return text
        self.cache[text] = out
        return out


def translate_data(data, tr):
    """description / reproduction / 발견 제목 / 위험평가근거(근거문)를 한국어로 번역."""
    for fd in data["findings"]:
        if fd.get("title"):
            fd["title"] = tr(fd["title"])
        if fd.get("description"):
            fd["description"] = tr(fd["description"])
        if fd.get("reproduction"):
            fd["reproduction"] = tr(fd["reproduction"])
        fd["risk_reasoning"] = [(m, s, tr(e)) for (m, s, e) in fd.get("risk_reasoning", [])]
    return data


# ══════════════════════════════════════════════════════════════════
#  2. DOCX 헬퍼
# ══════════════════════════════════════════════════════════════════
def cvss_to_isms_grade(score):
    try: s = float(score)
    except (TypeError, ValueError): return "(미상)"
    if s >= 9.0: return "상(긴급)"
    if s >= 7.0: return "상(높음)"
    if s >= 4.0: return "중"
    if s > 0.0:  return "하"
    return "정보"

def _set_eastasia(run):
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), FONT); rf.set(qn("w:ascii"), FONT); rf.set(qn("w:hAnsi"), FONT)

def add_run(p, text, bold=False, color="000000", size=10):
    r = p.add_run(str(text))
    r.font.name = FONT; r.font.bold = bold; r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    _set_eastasia(r)
    return r

def shade(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    for ex in tcpr.findall(qn("w:shd")):
        tcpr.remove(ex)
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), hex_color)
    valign = tcpr.find(qn("w:vAlign"))
    if valign is not None:
        valign.addprevious(sh)
    else:
        tcpr.append(sh)

def fill_cell(cell, text, bold=False, color="000000", size=10, bg=None, align=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    if align: p.alignment = align
    lines = text if isinstance(text, list) else [text]
    for i, ln in enumerate(lines):
        if i: p = cell.add_paragraph(); p.paragraph_format.space_after = Pt(1)
        add_run(p, ln, bold=bold, color=color, size=size)
    if bg: shade(cell, bg)

def manual_cell(cell, hint="〔작성 필요〕"):
    fill_cell(cell, hint, color="B45309", size=9, bg=FILL_MANUAL)

def make_table(doc, rows, cols, widths_mm):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"; t.autofit = False; t.allow_autofit = False
    # 고정 레이아웃 (tblLayout 은 tblCellMar/tblLook 보다 앞 — 스키마 순서)
    tblpr = t._tbl.tblPr
    for ex in tblpr.findall(qn("w:tblLayout")):
        tblpr.remove(ex)
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed")
    anchor = tblpr.find(qn("w:tblCellMar")) or tblpr.find(qn("w:tblLook"))
    if anchor is not None:
        anchor.addprevious(lay)
    else:
        tblpr.append(lay)
    for r in t.rows:
        for j, c in enumerate(r.cells):
            c.width = Mm(widths_mm[j])
    return t

def heading(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    add_run(p, text, bold=True, color="1F2937", size=13)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "1F2937")
    pbdr.append(bottom)
    pPr.insert(0, pbdr)   # pBdr는 spacing/ind 보다 앞 (스키마 순서)
    return p

def para(doc, runs, align=None, space_after=4):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
    if align: p.alignment = align
    for txt, kw in runs: add_run(p, txt, **kw)
    return p


# ══════════════════════════════════════════════════════════════════
#  3. DOCX 조립
# ══════════════════════════════════════════════════════════════════
def build_docx(data, out_path):
    doc = Document()
    st = doc.styles["Normal"]; st.font.name = FONT; st.font.size = Pt(10)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)   # A4
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Mm(20))
    W = 170  # 본문 폭(mm)

    f = data["findings"][0] if data["findings"] else {}

    # 제목
    para(doc, [("정보시스템 취약점 점검(모의침투) 결과보고서", dict(bold=True, size=17))],
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, [(f"대상 시스템: {data['target']}   |   점검 유형: 모의침투(Penetration Test)   |   "
                f"작성일: {data['tested_on']}", dict(size=9, color="555555"))],
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, [("■ 흰색 셀 = AWS Security Agent 리포트 자동 매핑   ■ 노랑 셀 = 〔작성 필요〕 조직 절차·승인 증적 수기 입력",
                dict(size=8, color="B45309"))], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # 1. 점검 개요
    heading(doc, "1. 점검 개요")
    overview = [
        ("점검 대상 URL", data["target_url"], False),
        ("점검 도구", "AWS AI Security Agent (자율 AI 모의침투)", False),
        ("점검 수행일", data["tested_on"], False),
        ("점검 방법", "AWS Security Agent 4단계 — ①Preflight(연결성) ②Static Analysis(코드·구성) "
                    "③Penetration Testing(런타임 익스플로잇) ④Finalizing(검증·리포팅). "
                    "각 발견은 proof-based exploitation으로 검증, CVSS v3.1 평가.", False),
        ("점검 수행 주체", "〔작성 필요〕 도구: AWS Security Agent / 검토·수행 책임자 명시(심사 책임소재 대응)", True),
        ("대외/내부 시스템 분류", "〔작성 필요〕 리포트 미제공 — 자산 분류로 부여(대외=반기, 내부=연1회)", True),
        ("점검 근거·주기 규정", "〔작성 필요〕 내부 규정상 점검 주기·근거(ISMS-P 2.11.2) 및 점검계획 승인 연계", True),
        ("데이터 처리·리전 유의", "〔검토 필요〕 Security Agent 처리/저장 리전 확인(서울 리전 부재) — 위·수탁/국외이전 검토", True),
    ]
    t = make_table(doc, len(overview), 2, [45, W - 45])
    for i, (k, v, manual) in enumerate(overview):
        fill_cell(t.rows[i].cells[0], k, bold=True, size=9, bg=FILL_LABEL)
        if manual: manual_cell(t.rows[i].cells[1], v)
        else: fill_cell(t.rows[i].cells[1], v, size=9)

    # 2. 리포트 범위 고지
    heading(doc, "2. 리포트 범위 고지 (적용 필터)")
    para(doc, [("주의: 본 리포트는 아래 필터가 적용된 ", dict(bold=True, size=9)),
               ("부분 뷰", dict(bold=True, size=9, color="C0392B")),
               ("이다. ISMS 정기점검 전체 증적으로 쓰려면 필터 미적용 전체 보고서를 별도 확보해야 한다.",
                dict(bold=True, size=9))])
    fl = data["filters"]
    t = make_table(doc, 2, 4, [28, 57, 28, W - 113])
    pairs = [("위험등급", fl["risk"], "신뢰도", fl["confidence"]),
             ("상태", fl["status"], "위험유형", fl["risktypes"])]
    for i, (a, b, c, d) in enumerate(pairs):
        fill_cell(t.rows[i].cells[0], a, bold=True, size=9, bg=FILL_LABEL)
        fill_cell(t.rows[i].cells[1], b, size=9)
        fill_cell(t.rows[i].cells[2], c, bold=True, size=9, bg=FILL_LABEL)
        fill_cell(t.rows[i].cells[3], d, size=9)

    # 3. 점검 결과 요약
    heading(doc, "3. 점검 결과 요약")
    c = data["counts"] or {}
    t = make_table(doc, 2, 6, [W/6]*6)
    hdr = ["총 발견", "긴급", "높음", "중간", "낮음", "정보"]
    vals = [data["total"], c.get("critical"), c.get("high"), c.get("medium"), c.get("low"), c.get("info")]
    for j, h in enumerate(hdr):
        fill_cell(t.rows[0].cells[j], h, bold=True, color="FFFFFF", size=9, bg=FILL_HEAD,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for j, v in enumerate(vals):
        em = (j in (0, 2)) and v
        fill_cell(t.rows[1].cells[j], "-" if v is None else str(v), bold=bool(em),
                  color="C0392B" if j == 2 and v else "000000", size=10,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, [("CVSS v3.1 → ISMS 위험등급 변환 기준(조직 정의 예시): ", dict(bold=True, size=8)),
               ("9.0+ 상(긴급) / 7.0–8.9 상(높음) / 4.0–6.9 중 / 0.1–3.9 하", dict(size=8, color="555555"))])

    # 4. 점검 범위 및 수행 항목
    heading(doc, "4. 점검 범위 및 수행 항목")
    tk = data["tasks"]
    para(doc, [(f"수행 위험유형 {len(tk['types'])}종 · 총 {tk['total']}개 태스크"
                f"(완료 {tk['completed']} · 중단 {tk['aborted']}) — 대외 공개 애플리케이션 전 엔드포인트 대상. 점검 범위 증적.",
                dict(size=9, color="555555"))])
    rows = tk["types"] or KNOWN_RISK_TYPES
    t = make_table(doc, len(rows) + 1, 2, [80, W - 80])
    fill_cell(t.rows[0].cells[0], "위험유형 (한글)", bold=True, color="FFFFFF", size=9, bg=FILL_HEAD)
    fill_cell(t.rows[0].cells[1], "Risk Type (원문)", bold=True, color="FFFFFF", size=9, bg=FILL_HEAD)
    for i, rt in enumerate(rows, 1):
        fill_cell(t.rows[i].cells[0], RISK_KR.get(rt, rt), bold=True, size=9)
        fill_cell(t.rows[i].cells[1], rt, size=9, color="555555")

    # 5. 발견사항 상세
    heading(doc, "5. 발견사항 상세")
    for idx, fd in enumerate(data["findings"], 1):
        para(doc, [(f"[발견 {idx}] {fd.get('title','')}", dict(bold=True, size=11, color="C0392B"))])
        grade = cvss_to_isms_grade(fd.get("score"))
        meta = make_table(doc, 5, 4, [30, 55, 28, W - 113])
        rows_meta = [
            ("발견 ID", fd.get("id", "-"), "위험등급", f"{fd.get('severity','-')} / {grade}"),
            ("CVSS 점수", fd.get("score", "-"), "신뢰도", fd.get("confidence", "-")),
            ("위험유형", fd.get("risktype", "-"), "상태", fd.get("status", "-")),
        ]
        for r, (a, b, cc, d) in enumerate(rows_meta):
            fill_cell(meta.rows[r].cells[0], a, bold=True, size=9, bg=FILL_LABEL)
            fill_cell(meta.rows[r].cells[1], b, size=8 if a == "발견 ID" else 9,
                      bold=(a == "CVSS 점수"))
            fill_cell(meta.rows[r].cells[2], cc, bold=True, size=9, bg=FILL_LABEL)
            fill_cell(meta.rows[r].cells[3], d, size=9,
                      bold=(cc == "위험등급"), color="C0392B" if cc == "위험등급" else "000000")
        # 식별 시각 / CVSS 벡터 (병합 행)
        fill_cell(meta.rows[3].cells[0], "식별 시각", bold=True, size=9, bg=FILL_LABEL)
        meta.rows[3].cells[1].merge(meta.rows[3].cells[3])
        fill_cell(meta.rows[3].cells[1], fd.get("identified", "-"), size=9)
        fill_cell(meta.rows[4].cells[0], "CVSS 벡터", bold=True, size=9, bg=FILL_LABEL)
        meta.rows[4].cells[1].merge(meta.rows[4].cells[3])
        fill_cell(meta.rows[4].cells[1], fd.get("cvss", "-"), size=8)

        if fd.get("description"):
            para(doc, [("취약점 설명", dict(bold=True, size=9))], space_after=2)
            para(doc, [(fd["description"], dict(size=9))])
        if fd.get("reproduction"):
            para(doc, [("재현 절차 (PoC)", dict(bold=True, size=9))], space_after=2)
            para(doc, [(fd["reproduction"][:1500], dict(size=8, color="333333"))])
        if fd.get("risk_reasoning"):
            para(doc, [("위험 평가 근거 (CVSS 메트릭)", dict(bold=True, size=9))], space_after=2)
            rr = fd["risk_reasoning"]
            rt = make_table(doc, len(rr) + 1, 2, [55, W - 55])
            fill_cell(rt.rows[0].cells[0], "메트릭", bold=True, color="FFFFFF", size=9, bg=FILL_HEAD)
            fill_cell(rt.rows[0].cells[1], "판정 및 근거", bold=True, color="FFFFFF", size=9, bg=FILL_HEAD)
            for r, (metric, short, expl) in enumerate(rr, 1):
                fill_cell(rt.rows[r].cells[0], METRIC_KR.get(metric, metric), bold=True, size=8, bg=FILL_LABEL)
                valv = VAL_KR.get(short.lower(), short)
                fill_cell(rt.rows[r].cells[1], f"{valv} — {expl}", size=8)
        para(doc, [("※ 위 권고/근거는 '발견·제안'이며 '조치 결과'가 아님 — 실제 조치·이행점검은 7장에서 수기 작성.",
                    dict(size=8, color="B45309"))])

    # 6. ISMS-P 인증기준 매핑
    heading(doc, "6. ISMS-P 인증기준 매핑")
    mapping = [("2.11.2 취약점 점검 및 조치", "정기 취약점 점검·이력관리·조치·이행점검", "직접 증적"),
               ("2.8.2 보안 요구사항 검토 및 시험", "배포 애플리케이션 보안 시험 수행", "보조 증적")]
    seen = {m[0] for m in mapping}
    for fd in data["findings"]:
        for row in RISK_TO_ISMS.get(fd.get("risktype", ""), []):
            if row[0] not in seen: mapping.append(row); seen.add(row[0])
    t = make_table(doc, len(mapping) + 1, 3, [58, W - 86, 28])
    for j, h in enumerate(["ISMS-P 인증기준", "연계 근거", "매핑 강도"]):
        fill_cell(t.rows[0].cells[j], h, bold=True, color="FFFFFF", size=9, bg=FILL_HEAD)
    for i, (cc, b, stg) in enumerate(mapping, 1):
        fill_cell(t.rows[i].cells[0], cc, bold=True, size=9)
        fill_cell(t.rows[i].cells[1], b, size=9)
        fill_cell(t.rows[i].cells[2], stg, size=9, bold=(stg == "직접 증적"),
                  color="1E7E34" if stg == "직접 증적" else "555555")

    # 7. 조치 및 이행점검
    heading(doc, "7. 조치 및 이행점검")
    t = make_table(doc, len(data["findings"]) + 1, 5, [W-4*26, 26, 26, 26, 26])
    for j, h in enumerate(["발견", "조치 담당자", "조치 기한", "조치 내역", "재점검 결과"]):
        fill_cell(t.rows[0].cells[j], h, bold=True, color="FFFFFF", size=9, bg=FILL_HEAD)
    for i, fd in enumerate(data["findings"], 1):
        fill_cell(t.rows[i].cells[0], fd.get("title", ""), size=8)
        for j in range(1, 5): manual_cell(t.rows[i].cells[j])
    para(doc, [("※ 이행점검 증적 = 조치 후 동일 PoC 재실행 리포트를 '조치 전/후 한 쌍'으로 첨부. "
                "자동 조치(PR) 사용 시 변경관리 승인 이력 별첨.", dict(size=8, color="B45309"))])

    # 8. 미조치 취약점 처리 및 보고·승인
    heading(doc, "8. 미조치 취약점 처리 및 보고·승인")
    t = make_table(doc, 4, 2, [45, W - 45])
    for i, (k, hint) in enumerate([
        ("미조치 사유", "〔작성 필요〕 즉시 조치 불가 시 사유 명확화"),
        ("잔여 위험·보완대책", "〔작성 필요〕 잔여위험 평가 및 보완대책"),
        ("책임자 보고 이력", "〔작성 필요〕 발견사항 책임자 보고 일시·방법"),
        ("책임자 승인", "〔작성 필요〕 성명 / 직책 / 승인일 / 서명")]):
        fill_cell(t.rows[i].cells[0], k, bold=True, size=9, bg=FILL_LABEL)
        manual_cell(t.rows[i].cells[1], hint)

    # 9. 결재
    heading(doc, "9. 결재")
    t = make_table(doc, 2, 3, [W/3]*3)
    for j, h in enumerate(["작성", "검토", "승인"]):
        fill_cell(t.rows[0].cells[j], h, bold=True, color="FFFFFF", size=9, bg=FILL_HEAD,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        manual_cell(t.rows[1].cells[j], "〔서명 / 일자〕")

    # settings.xml zoom percent 보정 (일부 Word 버전 호환)
    zoom = doc.settings.element.find(qn("w:zoom"))
    if zoom is not None and not zoom.get(qn("w:percent")):
        zoom.set(qn("w:percent"), "100")

    doc.save(out_path)
    return out_path


# ══════════════════════════════════════════════════════════════════
#  4. 진입점
# ══════════════════════════════════════════════════════════════════
def main():
    import argparse
    ap = argparse.ArgumentParser(
        description="AWS Security Agent 펜테스트 PDF → ISMS-P 점검 결과보고서(.docx) 생성")
    ap.add_argument("pdf", help="펜테스트 리포트 PDF 경로")
    ap.add_argument("out", nargs="?", help="출력 docx 경로(생략 시 자동 명명)")
    ap.add_argument("--no-translate", action="store_true",
                    help="한국어 번역 비활성화(영문 원문 유지)")
    ap.add_argument("--model", default=DEFAULT_MODEL_ID,
                    help=f"Bedrock 모델/추론 프로파일 ID (기본: {DEFAULT_MODEL_ID})")
    ap.add_argument("--region", default=os.environ.get("AWS_REGION") or os.environ.get("AWS_DEFAULT_REGION"),
                    help="Bedrock 리전 (기본: AWS_REGION 환경변수 / boto3 기본 세션)")
    args = ap.parse_args()

    if not os.path.isfile(args.pdf):
        print(f"파일을 찾을 수 없음: {args.pdf}"); sys.exit(1)

    data = extract_report(args.pdf)

    if not args.no_translate:
        tr = Translator(enabled=True, model_id=args.model, region=args.region)
        if tr.enabled:
            print(f"번역 중(한국어) … 모델={args.model}"
                  f"{' / 리전=' + args.region if args.region else ''}")
            translate_data(data, tr)

    out = args.out or f"ISMS-P_취약점점검결과보고서_{data['target']}.docx"
    build_docx(data, out)
    print(f"추출: 발견 {len(data['findings'])}건 / 위험유형 {len(data['tasks']['types'])}종 / "
          f"태스크 {data['tasks']['total']}개")
    print(f"생성 완료: {out}")


if __name__ == "__main__":
    main()
